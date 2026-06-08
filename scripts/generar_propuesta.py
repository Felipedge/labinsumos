from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# --- Styles helpers ---
def set_font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, color=(31, 73, 125)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=13 if level==1 else 11, bold=True, color=color)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(text, indent=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=10.5)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=10.5)
    p.paragraph_format.space_after = Pt(3)

def shade_row(row, hex_color='D6E4F0'):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

def cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# =====================================================================
# PORTADA
# =====================================================================
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PROPUESTA COMERCIAL')
set_font(run, size=22, bold=True, color=(31, 73, 125))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Plataforma LabInsumos')
set_font(run, size=16, bold=True, color=(68, 114, 196))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Software de Gestión de Insumos para Laboratorio Farmacéutico')
set_font(run, size=12, italic=True, color=(89, 89, 89))

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Modalidad SaaS · Arriendo mensual · Modelo por rango de usuarios')
set_font(run, size=10.5, color=(89, 89, 89))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
today = datetime.date.today()
meses = ['enero','febrero','marzo','abril','mayo','junio','julio','agosto','septiembre','octubre','noviembre','diciembre']
run = p.add_run(f'Junio de 2026')
set_font(run, size=10, color=(127, 127, 127))

doc.add_page_break()

# =====================================================================
# 1. RESUMEN EJECUTIVO
# =====================================================================
heading('1. Resumen Ejecutivo')
body(
    'LabInsumos es una plataforma web especializada en la gestión integral de insumos de laboratorio '
    'farmacéutico: estándares, reactivos, columnas, placebos y APIs. Desarrollada bajo tecnología '
    'moderna (React 18 + Firebase), cubre el ciclo completo de vida de cada insumo, desde la recepción '
    'y trazabilidad hasta la aprobación regulatoria y el reporte de auditoría.'
)
body(
    'La presente propuesta describe las condiciones de arriendo mensual bajo la modalidad Software as '
    'a Service (SaaS), expresadas en Unidades de Fomento (UF) para proteger tanto al proveedor como '
    'al cliente frente a variaciones inflacionarias.'
)

# =====================================================================
# 2. DESCRIPCIÓN DE LA PLATAFORMA
# =====================================================================
heading('2. Descripción de la Plataforma')

heading('2.1 Módulos incluidos', level=2)
modulos = [
    ('Dashboard',        'Vista general con semáforos de vencimiento y KPIs del inventario.'),
    ('Estándares',       'Registro, trazabilidad por frasco, historial de uso y alertas de vencimiento.'),
    ('Columnas HPLC',    'Control de ciclos, usos acumulados y estado operativo.'),
    ('Reactivos',        'Gestión de stock, vencimiento y consumo por análisis.'),
    ('Placebo',          'Trazabilidad de muestras de placebo por cliente y lote.'),
    ('APIs',             'Control de principios activos con stock en gramos y decimal de precisión.'),
    ('Alertas',          'Notificaciones automáticas de vencimiento próximo y crítico.'),
    ('Métodos',          'Vinculación de métodos analíticos con los insumos que requieren.'),
    ('Aprobaciones',     'Flujo de aprobación de insumos con roles diferenciados.'),
    ('Clientes',         'Gestión de clientes asociados a placebos y métodos.'),
    ('Usuarios',         'Administración de roles: Lectura, Administrativo, Analista, Coordinador, Jefe, Admin.'),
    ('Auditoría',        'Log completo de todas las operaciones realizadas en la plataforma.'),
    ('Reportes',         'Exportación de inventario y trazabilidad a Excel/PDF.'),
    ('Escaneo QR',       'Registro y consulta de insumos mediante código QR.'),
]
for nombre, desc in modulos:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    run1 = p.add_run(f'{nombre}: ')
    set_font(run1, size=10.5, bold=True)
    run2 = p.add_run(desc)
    set_font(run2, size=10.5)

heading('2.2 Características técnicas', level=2)
for item in [
    'Acceso web 100% desde navegador, sin instalación local.',
    'Base de datos en la nube (Google Firebase) con alta disponibilidad.',
    'Roles y permisos granulares por usuario.',
    'Interfaz responsive — compatible con tablet y escritorio.',
    'Actualizaciones automáticas incluidas en el arriendo.',
    'Backup automático gestionado por la infraestructura Firebase.',
]:
    bullet(item)

# =====================================================================
# 3. MODELO DE PRECIOS
# =====================================================================
heading('3. Modelo de Precios')
body(
    'El precio se expresa en Unidades de Fomento (UF), valor indexado al IPC publicado por el Banco '
    'Central de Chile. El cobro mensual se convierte a pesos chilenos (CLP) según el valor de la UF '
    'del día de facturación.'
)

heading('3.1 Tabla de precios por rango de usuarios activos', level=2)
body(
    'Se entiende por "usuario activo" todo usuario con acceso habilitado a la plataforma durante '
    'el período facturado, independientemente de la frecuencia de uso.'
)

# Table
tbl = doc.add_table(rows=1, cols=5)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'

hdrs = ['Rango de usuarios', 'Plan', 'UF / mes', 'CLP aprox. / mes*', 'USD aprox. / mes**']
for i, h in enumerate(hdrs):
    cell_text(tbl.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
shade_row(tbl.rows[0], '1F497D')
for cell in tbl.rows[0].cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

planes = [
    ('1 – 10',   'Básico',      '1,5 – 2,0',  '$58.000 – $77.000',   '$61 – $82'),
    ('11 – 20',  'Estándar',    '2,5 – 3,0',  '$96.000 – $116.000',  '$102 – $123'),
    ('21 – 30',  'Profesional', '3,5 – 4,5',  '$135.000 – $173.000', '$143 – $184'),
    ('31 – 40',  'Avanzado',    '5,0 – 6,0',  '$193.000 – $231.000', '$205 – $245'),
    ('41 – 50',  'Empresarial', '6,5 – 8,0',  '$250.000 – $308.000', '$266 – $328'),
    ('51 – 75',  'Corporativo', '9,0 – 11,0', '$347.000 – $424.000', '$369 – $450'),
    ('76 – 100', 'Enterprise',  '12,0 – 15,0','$462.000 – $578.000', '$491 – $614'),
]
colors = ['FFFFFF', 'EBF3FB']
for idx, (rango, plan, uf, clp, usd) in enumerate(planes):
    row = tbl.add_row()
    shade_row(row, colors[idx % 2])
    cell_text(row.cells[0], rango, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row.cells[1], plan,  align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row.cells[2], uf,    align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row.cells[3], clp,   align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row.cells[4], usd,   align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(
    '* CLP calculado con UF = $38.500 (referencial). '
    '** USD calculado con tipo de cambio $940/USD (referencial). '
    'Los valores definitivos se ajustan al día de facturación.'
)
set_font(run, size=9, italic=True, color=(127, 127, 127))

# =====================================================================
# 4. QUÉ INCLUYE EL ARRIENDO
# =====================================================================
heading('4. ¿Qué incluye el arriendo mensual?')

incluidos = {
    'Acceso a la plataforma':
        'Todos los módulos descritos en la sección 2, sin restricción de funcionalidades según el plan.',
    'Alojamiento e infraestructura':
        'Hospedaje en Google Firebase (Firestore, Hosting, Authentication). Escalado automático según demanda.',
    'Actualizaciones':
        'Mejoras de funcionalidad, correcciones de errores y parches de seguridad sin costo adicional.',
    'Soporte técnico':
        'Soporte por correo electrónico con tiempo de respuesta de 24 horas hábiles.',
    'Capacitación inicial':
        'Una sesión de onboarding de hasta 2 horas para el equipo administrador (remota vía videollamada).',
    'Respaldo de datos':
        'Backup automático gestionado por Firebase con retención de 7 días.',
}
for titulo, desc in incluidos.items():
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f'{titulo}: ')
    set_font(r1, size=10.5, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=10.5)

# =====================================================================
# 5. SERVICIOS OPCIONALES
# =====================================================================
heading('5. Servicios Opcionales (cotización separada)')

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl2.style = 'Table Grid'
for i, h in enumerate(['Servicio', 'Descripción', 'Precio referencial']):
    cell_text(tbl2.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
shade_row(tbl2.rows[0], '1F497D')
for cell in tbl2.rows[0].cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

opcionales = [
    ('Migración de datos',        'Importación de inventario histórico desde Excel u otro sistema.',      '3 – 8 UF (única vez)'),
    ('Onboarding extendido',      'Capacitación a usuarios finales, hasta 4 horas.',                      '2 – 4 UF (única vez)'),
    ('Desarrollo a medida',       'Nuevas funcionalidades o integraciones solicitadas por el cliente.',   'Cotización por hora'),
    ('Soporte prioritario',       'Tiempo de respuesta de 4 horas hábiles + línea WhatsApp Business.',   '0,5 – 1,0 UF/mes'),
    ('Usuarios adicionales >100', 'Incorporación de usuarios por sobre el rango Enterprise.',             '0,1 – 0,15 UF/usuario/mes'),
]
for idx, (srv, desc, precio) in enumerate(opcionales):
    row = tbl2.add_row()
    shade_row(row, 'FFFFFF' if idx%2==0 else 'EBF3FB')
    cell_text(row.cells[0], srv,    bold=True)
    cell_text(row.cells[1], desc)
    cell_text(row.cells[2], precio, align=WD_ALIGN_PARAGRAPH.CENTER)

# =====================================================================
# 6. CONDICIONES COMERCIALES
# =====================================================================
heading('6. Condiciones Comerciales')

condiciones = [
    ('Período mínimo de contrato', 'Tres (3) meses.'),
    ('Facturación',                'Mensual, anticipada, emitida los primeros 5 días hábiles de cada mes.'),
    ('Forma de pago',              'Transferencia bancaria o cheque a 30 días.'),
    ('Descuento por contrato anual', 'Pago adelantado de 12 meses: descuento del 12% sobre la tarifa mensual.'),
    ('Descuento por contrato bianual', 'Pago adelantado de 24 meses: descuento del 20% sobre la tarifa mensual.'),
    ('Ajuste de tarifa',           'Revisión anual en función de la variación acumulada del IPC. La UF ya '
                                   'absorbe parte de esta variación; el ajuste aplica sólo al número de UF pactado.'),
    ('Suspensión del servicio',    'Ante mora superior a 30 días corridos, se suspende el acceso hasta regularización.'),
    ('Rescisión',                  'Aviso con 30 días de anticipación. No aplica cobro de penalidad.'),
    ('Confidencialidad',           'Los datos del cliente son de su exclusiva propiedad. '
                                   'El proveedor no cede ni comercializa información del cliente.'),
    ('Propiedad intelectual',      'La plataforma es propiedad del proveedor. '
                                   'El cliente recibe una licencia de uso no exclusiva durante la vigencia del contrato.'),
]
for titulo, desc in condiciones:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f'{titulo}: ')
    set_font(r1, size=10.5, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=10.5)

# =====================================================================
# 7. PROCESO DE IMPLEMENTACIÓN
# =====================================================================
heading('7. Proceso de Implementación')

etapas = [
    ('Semana 1', 'Firma de contrato y pago del primer mes. Creación del entorno de producción y configuración inicial.'),
    ('Semana 1–2', 'Migración de datos históricos (si aplica) y configuración de usuarios con sus roles correspondientes.'),
    ('Semana 2', 'Sesión de onboarding para el equipo administrador y coordinadores.'),
    ('Semana 3–4', 'Operación asistida: período de acompañamiento con soporte prioritario sin costo adicional.'),
    ('Desde el mes 2', 'Operación autónoma con soporte estándar incluido.'),
]
for etapa, desc in etapas:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f'{etapa}: ')
    set_font(r1, size=10.5, bold=True, color=(31, 73, 125))
    r2 = p.add_run(desc)
    set_font(r2, size=10.5)

# =====================================================================
# 8. CUMPLIMIENTO REGULATORIO
# =====================================================================
heading('8. Cumplimiento Regulatorio')
body(
    'LabInsumos está diseñado alineado con los requerimientos habituales de laboratorios farmacéuticos '
    'bajo normativas como BPL (Buenas Prácticas de Laboratorio) y los estándares de la ICH Q2(R1) '
    'para validación analítica. La plataforma provee:'
)
for item in [
    'Trazabilidad completa de cada insumo desde su ingreso hasta su baja o vencimiento.',
    'Log de auditoría inalterable con registro de usuario, fecha/hora y operación realizada.',
    'Flujo de aprobación multinivel para el ingreso de insumos críticos.',
    'Control de vencimientos con alertas visuales de tres niveles (OK / Advertencia / Crítico).',
    'Histórico de uso por método analítico, analista y cliente.',
]:
    bullet(item)

body(
    'Nota: el proveedor no certifica el cumplimiento regulatorio del laboratorio contratante; '
    'la plataforma es una herramienta de apoyo a la gestión. La validación formal de los procesos '
    'documentales es responsabilidad del equipo de Calidad del laboratorio.',
    space_after=4
)

# =====================================================================
# 9. PRÓXIMOS PASOS
# =====================================================================
heading('9. Próximos Pasos')
for step in [
    'Revisión de esta propuesta por parte del cliente.',
    'Reunión de alineamiento para ajustar plan, módulos y condiciones (opcional).',
    'Firma del contrato de arriendo.',
    'Pago del primer período y activación del entorno de producción.',
]:
    bullet(step)

# =====================================================================
# CIERRE
# =====================================================================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Propuesta válida por 30 días corridos desde su emisión —')
set_font(run, size=9.5, italic=True, color=(127, 127, 127))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Para consultas o ajustes a esta propuesta, contactar al equipo comercial.')
set_font(run, size=10)

# Save
out = '/home/user/labinsumos/Propuesta_LabInsumos_SaaS.docx'
doc.save(out)
print(f'Guardado en: {out}')
